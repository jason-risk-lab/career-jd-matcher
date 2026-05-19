"""Build resume/cover-letter/analysis-report .docx files from a JSON manifest."""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


KOREAN_FONT_PRIMARY = "Malgun Gothic"
KOREAN_FONT_FALLBACK = "Noto Sans CJK KR"
MONO_FONT = "Consolas"

HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 11}
BODY_SIZE = 10.5


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, mono=False, color=None):
    """Apply Korean-aware fonts and styling to a run."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    name = MONO_FONT if mono else KOREAN_FONT_PRIMARY
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def set_page(doc):
    """Configure A4 size and 25mm margins."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)


def set_default_style(doc):
    """Apply Korean font to the Normal style as a baseline."""
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT_PRIMARY
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), KOREAN_FONT_PRIMARY)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(paragraph, text, base_size=BODY_SIZE, base_bold=False, base_color=None):
    """Tokenize **bold**, *italic*, `code` and append runs."""
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_run_font(run, size=base_size, bold=base_bold, color=base_color)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, bold=base_bold, italic=True, color=base_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, bold=base_bold, mono=True, color=base_color)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, bold=base_bold, color=base_color)


def parse_markdown(text):
    """Parse markdown into a flat list of block tokens."""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            blocks.append(("code", "\n".join(buf)))
            i = j + 1
            continue
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            blocks.append(("hr", None))
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            blocks.append(("h" + str(len(m.group(1))), m.group(2).strip()))
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            blocks.append(("table", tbl))
            continue
        bm = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if bm:
            indent = len(bm.group(1))
            level = 2 if indent >= 2 else 1
            blocks.append(("ul", level, bm.group(3).strip()))
            i += 1
            continue
        nm = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if nm:
            indent = len(nm.group(1))
            level = 2 if indent >= 2 else 1
            blocks.append(("ol", level, nm.group(3).strip()))
            i += 1
            continue
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r"^#{1,4}\s", nxt) or nxt.startswith("```") or nxt.startswith("|"):
                break
            if re.match(r"^[-*]\s", nxt) or re.match(r"^\d+\.\s", nxt):
                break
            if re.match(r"^-{3,}$", nxt):
                break
            buf.append(nxt)
            i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks


def parse_table_rows(raw_rows):
    """Split pipe-delimited rows, drop the alignment row, return alignments + cells."""
    def split_row(row):
        cells = row.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    rows = [split_row(r) for r in raw_rows]
    alignments = None
    if len(rows) >= 2:
        sep = rows[1]
        if all(re.match(r"^:?-{2,}:?$", c) for c in sep if c):
            alignments = []
            for c in sep:
                left = c.startswith(":")
                right = c.endswith(":")
                if left and right:
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                elif right:
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            rows = [rows[0]] + rows[2:]
    return rows, alignments


def add_table(doc, raw_rows):
    """Render a pipe-syntax markdown table into the document."""
    rows, alignments = parse_table_rows(raw_rows)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            if alignments and c_idx < len(alignments):
                para.alignment = alignments[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            add_inline(para, text, base_bold=(r_idx == 0))


def add_hr(doc, page_break=False):
    """Add a horizontal rule, optionally as a page break."""
    if page_break:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        return
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(doc, level, text):
    """Append a styled heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    size = HEADING_SIZES.get(level, BODY_SIZE)
    color = RGBColor(0x1F, 0x3A, 0x5F) if level == 1 else RGBColor(0x2C, 0x3E, 0x50)
    set_run_font(run, size=size, bold=True, color=color)


def add_paragraph(doc, text):
    """Append a body paragraph with inline formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, text)


def add_list_item(doc, level, text, ordered=False, index=1):
    """Append a bullet or numbered list item with manual indentation."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6 * level)
    p.paragraph_format.space_after = Pt(2)
    prefix = f"{index}. " if ordered else ("  • " if level == 2 else "• ")
    run = p.add_run(prefix)
    set_run_font(run)
    add_inline(p, text)


def add_code_block(doc, code):
    """Append a monospaced paragraph for fenced code."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    set_run_font(run, mono=True, size=BODY_SIZE - 1)


def render_blocks(doc, blocks, hr_as_page_break=False):
    """Walk parsed markdown blocks and emit docx elements."""
    ol_counter = {1: 0, 2: 0}
    prev_kind = None
    for block in blocks:
        kind = block[0]
        if kind == "ol":
            level = block[1]
            ol_counter[level] += 1
            if level == 1:
                ol_counter[2] = 0
            add_list_item(doc, level, block[2], ordered=True, index=ol_counter[level])
        else:
            if prev_kind == "ol" and kind != "ol":
                ol_counter = {1: 0, 2: 0}
            if kind in ("h1", "h2", "h3", "h4"):
                add_heading(doc, int(kind[1]), block[1])
            elif kind == "p":
                add_paragraph(doc, block[1])
            elif kind == "ul":
                add_list_item(doc, block[1], block[2], ordered=False)
            elif kind == "table":
                add_table(doc, block[1])
            elif kind == "code":
                add_code_block(doc, block[1])
            elif kind == "hr":
                add_hr(doc, page_break=hr_as_page_break)
        prev_kind = kind


def add_page_number_field(paragraph, kind):
    """Insert a Word PAGE or NUMPAGES field."""
    run = paragraph.add_run()
    set_run_font(run, size=9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE " if kind == "PAGE" else " NUMPAGES "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


def add_header_footer(doc, company_name):
    """Add header with company report title and footer with Page X / Y."""
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(f"[{company_name}] 지원 분석 보고서")
    set_run_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Page ")
    set_run_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))
    add_page_number_field(footer_p, "PAGE")
    run = footer_p.add_run(" / ")
    set_run_font(run, size=9, color=RGBColor(0x55, 0x55, 0x55))
    add_page_number_field(footer_p, "NUMPAGES")


def add_cover_page(doc, company_name, job_title, meta, analysis_date):
    """Render the analysis report cover page."""
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name} {job_title}")
    set_run_font(run, size=24, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("지원 분석 보고서")
    set_run_font(run, size=18, bold=True, color=RGBColor(0x2C, 0x3E, 0x50))

    doc.add_paragraph()
    doc.add_paragraph()

    if meta:
        score = meta.get("total_score")
        if score is not None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"종합 점수: {score}")
            set_run_font(run, size=16, bold=True)
        rec = meta.get("recommendation")
        if rec:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(rec)
            set_run_font(run, size=14, bold=True, color=RGBColor(0x16, 0xA0, 0x85))
        conclusion = meta.get("one_line_conclusion")
        if conclusion:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(conclusion)
            set_run_font(run, size=12, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    for _ in range(4):
        doc.add_paragraph()

    if analysis_date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"분석일: {analysis_date}")
        set_run_font(run, size=11, color=RGBColor(0x55, 0x55, 0x55))

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def build_document(markdown_text, output_path, *, is_report=False, company_name="",
                   job_title="", analysis_date="", cover_meta=None):
    """Build one .docx file end to end."""
    doc = Document()
    set_page(doc)
    set_default_style(doc)
    if is_report:
        add_header_footer(doc, company_name)
        add_cover_page(doc, company_name, job_title, cover_meta or {}, analysis_date)
    blocks = parse_markdown(markdown_text)
    render_blocks(doc, blocks, hr_as_page_break=is_report)
    doc.save(str(output_path))


def process_manifest(manifest):
    """Process all files described in the manifest and report results."""
    output_dir = Path(manifest["output_dir"])
    output_dir.mkdir(parents=True, exist_ok=True)
    company = manifest.get("company_name", "")
    job_title = manifest.get("job_title", "")
    analysis_date = manifest.get("analysis_date", "")
    files = manifest.get("files", {})

    failed = False
    for key in ("resume", "cover_letter", "analysis_report"):
        entry = files.get(key)
        if not entry:
            continue
        filename = entry.get("filename", f"{key}.docx")
        try:
            md_path = Path(entry["markdown_path"])
            md_text = md_path.read_text(encoding="utf-8")
            out_path = output_dir / filename
            build_document(
                md_text,
                out_path,
                is_report=(key == "analysis_report"),
                company_name=company,
                job_title=job_title,
                analysis_date=analysis_date,
                cover_meta=entry.get("cover_meta"),
            )
            print(f"OK {filename}")
        except Exception as exc:
            failed = True
            print(f"ERROR {filename}: {exc}")
    return 1 if failed else 0


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(description="Build career-jd-matcher .docx outputs.")
    parser.add_argument("manifest", help="Path to JSON manifest file.")
    args = parser.parse_args()
    manifest = json.loads(Path(args.manifest).read_text(encoding="utf-8"))
    sys.exit(process_manifest(manifest))


if __name__ == "__main__":
    main()


# --- USAGE EXAMPLE ---
"""
Run:
    python build_docx.py manifest.json

Sample manifest.json:
{
  "output_dir": "/mnt/user-data/outputs/",
  "company_name": "루닛",
  "job_title": "내부감사 책임자",
  "analysis_date": "2026-05-15",
  "files": {
    "resume": {
      "filename": "경력기술서_루닛_내부감사_정교화.docx",
      "markdown_path": "/tmp/resume_refined.md"
    },
    "cover_letter": {
      "filename": "자기소개서_루닛_내부감사_정교화.docx",
      "markdown_path": "/tmp/cover_letter_refined.md"
    },
    "analysis_report": {
      "filename": "지원분석보고서_루닛_내부감사.docx",
      "markdown_path": "/tmp/analysis_report.md",
      "cover_meta": {
        "total_score": 71.8,
        "recommendation": "🟢 강력 지원",
        "one_line_conclusion": "성장성·정체성 일치 양호, Safety 축에 면접 시 확인 필요."
      }
    }
  }
}

Stdout will contain one line per file:
    OK 경력기술서_루닛_내부감사_정교화.docx
    OK 자기소개서_루닛_내부감사_정교화.docx
    OK 지원분석보고서_루닛_내부감사.docx

Exit code 0 on full success, 1 if any file failed.
"""
